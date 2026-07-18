"""
Módulo de extracción de texto.

Cada documento fuente (PDF, Word, Markdown, y en el futuro Excel/PowerPoint/CSV/JSON/HTML)
requiere una estrategia distinta de extracción, como se describe en la tarjeta
"2 - Proceso y extracción de contenido" del Trello del reto.

Este módulo centraliza esa lógica: dado un path de archivo, devuelve texto plano limpio
listo para pasar a la etapa de chunking.
"""

import os
import re
from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument
import markdown as md_lib
from bs4 import BeautifulSoup


def extract_from_pdf(filepath: str) -> str:
    """Extrae texto de un PDF nativo (no escaneado). Concatena el texto de todas las páginas."""
    reader = PdfReader(filepath)
    pages_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages_text.append(text)
    return "\n\n".join(pages_text)


def extract_from_docx(filepath: str) -> str:
    """Extrae texto de un documento Word, preservando párrafos y encabezados como líneas separadas."""
    doc = DocxDocument(filepath)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # También extraemos texto de tablas, si las hay (ver "Guía de Envíos")
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n\n".join(parts)


def extract_from_markdown(filepath: str) -> str:
    """Convierte Markdown a texto plano, eliminando la sintaxis pero conservando el contenido legible."""
    with open(filepath, "r", encoding="utf-8") as f:
        raw_md = f.read()
    html = md_lib.markdown(raw_md)
    soup = BeautifulSoup(html, "html.parser")
    text = soup.get_text(separator="\n")
    return text


def clean_text(text: str) -> str:
    """Limpieza básica: espacios duplicados, líneas vacías repetidas, saltos de línea excesivos."""
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    return text.strip()


# Mapea extensión de archivo -> función de extracción
EXTRACTORS = {
    ".pdf": extract_from_pdf,
    ".docx": extract_from_docx,
    ".md": extract_from_markdown,
}


def extract_text(filepath: str) -> str:
    """
    Punto de entrada único: detecta el formato por extensión y aplica el extractor correcto.
    Lanza ValueError si el formato no está soportado todavía.
    """
    ext = Path(filepath).suffix.lower()
    extractor = EXTRACTORS.get(ext)
    if extractor is None:
        raise ValueError(
            f"Formato no soportado: '{ext}'. Soportados actualmente: {list(EXTRACTORS.keys())}"
        )
    raw_text = extractor(filepath)
    return clean_text(raw_text)


if __name__ == "__main__":
    # Prueba rápida: extrae texto de todos los documentos en docs/
    docs_dir = Path(__file__).parent.parent / "docs"
    for filepath in sorted(docs_dir.glob("*")):
        if filepath.suffix.lower() in EXTRACTORS:
            print(f"\n{'=' * 60}\n{filepath.name}\n{'=' * 60}")
            text = extract_text(str(filepath))
            print(text[:300] + ("..." if len(text) > 300 else ""))
            print(f"\n[{len(text)} caracteres extraídos]")
